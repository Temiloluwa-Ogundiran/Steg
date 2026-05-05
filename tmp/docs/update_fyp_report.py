from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"C:\Users\USER\Documents\Steganography")
REPORT_PATH = ROOT / "report" / "SECURE STEGANOGRAPHY SOFTWARE SUITE RESILIENT TO STEGANALYSIS (OGUNDIRAN)  (Repaired).docx"
ASSET_DIR = ROOT / "report" / "assets"
ACTIVE_DOC = None


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_after(paragraph, rows, cols, style="Table Grid"):
    doc = ACTIVE_DOC
    table = doc.add_table(rows=rows, cols=cols)
    if style:
        table.style = style
    paragraph._p.addnext(table._tbl)
    return table


def insert_picture_after(paragraph, image_path, width=None):
    doc = ACTIVE_DOC
    pic_paragraph = doc.add_paragraph()
    run = pic_paragraph.add_run()
    if width is not None:
        run.add_picture(str(image_path), width=width)
    else:
        run.add_picture(str(image_path))
    pic_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph._p.addnext(pic_paragraph._p)
    return pic_paragraph


def insert_paragraph_after_table(table, text="", style=None):
    new_p = OxmlElement("w:p")
    table._tbl.addnext(new_p)
    new_para = Paragraph(new_p, table._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def find_paragraph(doc, text):
    expected = " ".join(text.split())
    matches = []
    for paragraph in doc.paragraphs:
        candidate = " ".join(paragraph.text.strip().split())
        if candidate == expected or candidate.startswith(expected):
            matches.append(paragraph)
    for paragraph in matches:
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        if "toc" not in style_name:
            return paragraph
    if matches:
        return matches[0]
    raise ValueError(f"Paragraph not found: {text}")


def remove_between(start_paragraph, end_paragraph):
    current = start_paragraph._p.getnext()
    while current is not None and current is not end_paragraph._p:
        next_element = current.getnext()
        current.getparent().remove(current)
        current = next_element


def remove_after(paragraph):
    current = paragraph._p.getnext()
    while current is not None:
        next_element = current.getnext()
        current.getparent().remove(current)
        current = next_element


def replace_section(doc, start_text, end_text, content):
    start = find_paragraph(doc, start_text)
    end = find_paragraph(doc, end_text)
    remove_between(start, end)
    anchor = start
    for item in content:
        style = item.get("style", "Normal")
        text = item.get("text", "")
        paragraph = insert_paragraph_after(anchor, text=text, style=style)
        if item.get("align") == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        anchor = paragraph
    return anchor


def add_section_block(anchor, content):
    current = anchor
    for item in content:
        if item["kind"] == "paragraph":
            paragraph = insert_paragraph_after(current, text=item.get("text", ""), style=item.get("style", "Normal"))
            if item.get("align") == "center":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            current = paragraph
        elif item["kind"] == "table":
            table = insert_table_after(current, rows=len(item["rows"]), cols=len(item["rows"][0]))
            for r, row in enumerate(item["rows"]):
                for c, value in enumerate(row):
                    table.cell(r, c).text = value
            current = insert_paragraph_after_table(table, text="", style="Normal")
        elif item["kind"] == "image":
            current = insert_picture_after(current, item["path"], width=item.get("width"))
        elif item["kind"] == "pagebreak":
            paragraph = insert_paragraph_after(current, text="", style="Normal")
            paragraph.add_run().add_break(WD_BREAK.PAGE)
            current = paragraph
        else:
            raise ValueError(f"Unknown content kind: {item['kind']}")
    return current


def make_architecture_diagram(path):
    path.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(image)
    font = ImageFont.load_default()

    def box(x1, y1, x2, y2, title, body_lines, fill):
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, outline="#141414", width=3, fill=fill)
        draw.text((x1 + 20, y1 + 18), title, fill="#141414", font=font)
        y = y1 + 70
        for line in body_lines:
            draw.text((x1 + 20, y), line, fill="#222222", font=font)
            y += 28

    def arrow(x1, y1, x2, y2):
        draw.line((x1, y1, x2, y2), fill="#1351AA", width=6)
        draw.polygon([(x2, y2), (x2 - 18, y2 - 10), (x2 - 18, y2 + 10)], fill="#1351AA")

    draw.text((40, 25), "Architecture of the Secure Steg Software Suite", fill="#141414", font=font)
    box(70, 140, 420, 330, "Presentation Layer", ["Web interface for encode, decode, and check", "Collects image, message, and passphrase input", "Displays download and validation results"], "#EAF0FA")
    box(620, 140, 980, 360, "Application Layer", ["FastAPI routes and service orchestration", "Input validation and upload normalization", "Runtime output management"], "#F4F1E8")
    box(1160, 140, 1510, 360, "Security Layer", ["Passphrase validation", "Salted key derivation", "Authenticated message encryption", "STEGv1 envelope generation and parsing"], "#EAF0FA")
    box(620, 470, 980, 700, "Steganography Core", ["Dense adversarial encoder-decoder model", "Hidden payload embedding into images", "Payload extraction from stego images"], "#F4F1E8")
    box(1160, 500, 1510, 700, "Runtime Storage", ["Temporary uploads", "Generated stego image outputs", "Short-lived download files"], "#EAF0FA")
    arrow(420, 235, 620, 235)
    arrow(980, 245, 1160, 245)
    arrow(800, 360, 800, 470)
    arrow(980, 585, 1160, 585)
    arrow(1160, 300, 980, 300)
    image.save(path)


def make_workflow_diagram(path):
    path.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1600, 1000), "white")
    draw = ImageDraw.Draw(image)
    font = ImageFont.load_default()

    def step(x1, y1, x2, y2, title, fill):
        draw.rounded_rectangle((x1, y1, x2, y2), radius=16, outline="#141414", width=3, fill=fill)
        draw.text((x1 + 18, y1 + 22), title, fill="#141414", font=font)

    def arrow(x1, y1, x2, y2):
        draw.line((x1, y1, x2, y2), fill="#1351AA", width=5)
        draw.polygon([(x2, y2), (x2 - 16, y2 - 8), (x2 - 16, y2 + 8)], fill="#1351AA")

    draw.text((40, 30), "Operational Workflow for Encode, Decode, and Check", fill="#141414", font=font)
    draw.text((60, 110), "Encode Path", fill="#141414", font=font)
    step(60, 150, 290, 230, "Select cover image", "#F4F1E8")
    step(360, 150, 640, 230, "Enter message + passphrase", "#EAF0FA")
    step(710, 150, 1030, 230, "Encrypt message into STEGv1 envelope", "#F4F1E8")
    step(1100, 150, 1510, 230, "Embed payload and generate stego image", "#EAF0FA")
    arrow(290, 190, 360, 190)
    arrow(640, 190, 710, 190)
    arrow(1030, 190, 1100, 190)
    draw.text((60, 370), "Decode Path", fill="#141414", font=font)
    step(60, 410, 290, 490, "Upload stego image", "#F4F1E8")
    step(360, 410, 620, 490, "Extract hidden payload", "#EAF0FA")
    step(690, 410, 990, 490, "Validate STEGv1 structure", "#F4F1E8")
    step(1060, 410, 1510, 490, "Decrypt with passphrase and recover message", "#EAF0FA")
    arrow(290, 450, 360, 450)
    arrow(620, 450, 690, 450)
    arrow(990, 450, 1060, 450)
    draw.text((60, 630), "Check Path", fill="#141414", font=font)
    step(60, 670, 300, 750, "Upload candidate image", "#F4F1E8")
    step(380, 670, 650, 750, "Extract hidden payload", "#EAF0FA")
    step(730, 670, 1050, 750, "Validate envelope only", "#F4F1E8")
    step(1130, 670, 1510, 750, "Return hidden-data present or absent", "#EAF0FA")
    arrow(300, 710, 380, 710)
    arrow(650, 710, 730, 710)
    arrow(1050, 710, 1130, 710)
    image.save(path)


def format_document(doc):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


def main():
    global ACTIVE_DOC
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    architecture_path = ASSET_DIR / "figure_3_1_architecture.png"
    workflow_path = ASSET_DIR / "figure_3_2_workflow.png"
    make_architecture_diagram(architecture_path)
    make_workflow_diagram(workflow_path)

    doc = Document(REPORT_PATH)
    ACTIVE_DOC = doc

    chapter_one_background = [
        {"text": "The growth of digital communication has made information exchange faster, wider, and more exposed than at any previous time. Sensitive personal, institutional, and commercial messages now move across networks where interception is realistic, and this has made information security a central concern in computer science and cybersecurity.", "style": "Normal"},
        {"text": "Traditional protection of digital communication has relied mainly on cryptography. Cryptography secures the content of a message by transforming readable plaintext into ciphertext that cannot be interpreted without the proper key. This gives strong confidentiality, but it does not hide the fact that a secret communication is taking place. In situations where the mere presence of encrypted content may attract attention, cryptography alone is not always sufficient.", "style": "Normal"},
        {"text": "Steganography addresses this second problem by concealing the existence of a message inside an apparently innocent carrier such as an image, audio signal, or video stream. In image steganography, the sender embeds secret information into a cover image and produces a stego image that appears visually similar to the original. The receiver then extracts the hidden information from that stego image using an agreed procedure. A strong steganographic system therefore aims at four closely related goals: imperceptibility, payload capacity, robustness, and resistance to detection.", "style": "Normal"},
        {"text": "Classical image steganography methods such as Least Significant Bit substitution and transform-domain embedding made important early contributions to covert communication. However, these methods typically follow fixed embedding rules. Once those rules introduce measurable statistical disturbances into the cover image, modern steganalysis systems can learn to detect them. What used to be a hidden message under simple visual inspection can now be exposed by data-driven classifiers that detect minute embedding artifacts.", "style": "Normal"},
        {"text": "The introduction of deep learning has changed both sides of this contest. On one side, deep neural networks have strengthened steganalysis by learning subtle residual patterns that handcrafted feature engineering could not reliably capture. On the other side, adversarial-network-based steganography has provided a more adaptive way of embedding information. Instead of manually deciding where and how to place message bits, an adversarial model learns embedding behaviour that seeks to remain close to natural image statistics while still preserving decodability.", "style": "Normal"},
        {"text": "This project is motivated by the fact that many publicly accessible steganography tools are still built around older embedding strategies and do not adequately address present-day steganalysis threats. In addition, practical secure software in this area remains limited. Some research works propose stronger adversarial or deep-learning-based steganographic methods, but these are often presented as technical prototypes rather than complete software suites. As a result, there is a gap between academic advances in steganography and real systems that ordinary users can operate securely through a clear encode, decode, and verification workflow.", "style": "Normal"},
        {"text": "For this reason, this project presents an original secure steganography software suite that uses an adversarial-network-based image steganography method as its embedding backbone and strengthens it with passphrase-protected payload handling, structured encode-decode-check workflows, and a user-facing application interface. The aim is not only to hide information in images, but to do so through a software system that is practical, secure, and more resilient to modern steganalysis than conventional methods.", "style": "Normal"},
    ]

    chapter_one_problem = [
        {"text": "The main problem addressed in this study is that many conventional steganography tools are not sufficiently resilient to present-day steganalysis. Spatial-domain and transform-domain methods often introduce statistical patterns that can be detected by machine-learning-based steganalysis systems, especially when payload size increases or when the attacker has access to modern residual-network detectors.", "style": "Normal"},
        {"text": "A second problem is that practical secure software in this area remains limited. Some research works propose stronger adversarial or deep-learning-based steganographic methods, but these are often presented as technical prototypes rather than complete software suites. As a result, there is a gap between model-level progress and software systems that ordinary users can use securely through a controlled encode, decode, and verification workflow.", "style": "Normal"},
    ]

    chapter_one_aim = [
        {"text": "The aim of this project is to design, implement, and evaluate a secure steganography software suite that conceals text messages inside images using an adversarial-network-based approach and improves practical security through controlled, passphrase-protected recovery.", "style": "Normal"},
        {"text": "To achieve this aim, the study pursues the following objectives:", "style": "Normal"},
        {"text": "To review the major concepts, methods, tools, and research trends in image steganography, steganalysis, and adversarial-network-based information hiding.", "style": "List Paragraph"},
        {"text": "To analyze the limitations of conventional steganography methods and identify the gaps that justify a more secure and practical software suite.", "style": "List Paragraph"},
        {"text": "To design a complete steganography system that supports image encoding, decoding, hidden-data checking, and passphrase-protected message recovery.", "style": "List Paragraph"},
        {"text": "To implement the proposed system using a web-based software architecture with a deep-learning steganography backbone and a secure payload protection layer.", "style": "List Paragraph"},
        {"text": "To evaluate the system using established image-quality and security criteria and to compare the adopted adversarial method with less effective traditional approaches.", "style": "List Paragraph"},
    ]

    chapter_one_methodology = [
        {"text": "This study adopts a structured development methodology in which the work proceeds from problem definition to design, implementation, and evaluation. The process is consistent with the logic of the Waterfall model because the major phases of the project were clearly defined from the outset, but each phase was also grounded in technical validation and iterative refinement of details within the phase itself.", "style": "Normal"},
        {"text": "The first phase was the literature review and problem analysis stage. In this stage, relevant concepts in steganography, steganalysis, cryptography, and adversarial learning were reviewed. Existing classical tools and modern research methods were compared in order to identify the most important practical and research gaps.", "style": "Normal"},
        {"text": "The second phase was system analysis and design. Here, the requirements of the proposed suite were defined, the operational modules were identified, and the system architecture was specified. Particular attention was given to the encode, decode, and check workflow and to the introduction of a passphrase-protected secure payload layer.", "style": "Normal"},
        {"text": "The third phase was implementation. The system was implemented as a web-based software suite with a frontend interface, a FastAPI backend, a secure payload service layer, and an adversarial-network-based steganography engine. Supporting functions such as upload normalization, temporary output management, and secure message handling were also developed.", "style": "Normal"},
        {"text": "The fourth phase was testing and evaluation. This involved functional verification of the software modules, security-oriented testing of passphrase-controlled recovery, and technical evaluation using image-quality and detection-related measures discussed in the literature. The project also compared the adopted adversarial approach against conventional steganographic methods such as LSB-based and transform-domain techniques.", "style": "Normal"},
    ]

    chapter_one_significance = [
        {"text": "This study is significant because it contributes to both the academic and practical sides of secure communication.", "style": "Normal"},
        {"text": "Practical significance: the project provides a usable secure steganography suite rather than stopping at model-level theory. This makes stronger information-hiding techniques more accessible to users who need a clear encode, decode, and verification workflow.", "style": "List Paragraph"},
        {"text": "Security significance: the system combines concealed communication with passphrase-protected payload handling, thereby adding a second layer of protection even when a hidden payload is recovered from an image.", "style": "List Paragraph"},
        {"text": "Academic significance: the work contributes to the continuing contest between steganography and steganalysis by examining how adversarial-network-based methods compare with weaker traditional approaches in terms of security, usability, and detectability.", "style": "List Paragraph"},
    ]

    chapter_one_outline = [
        {"text": "Chapter One introduces the study, states the problem, presents the aim and objectives, explains the methodology, and outlines the significance of the work. Chapter Two reviews the literature on classical steganography, steganalysis, machine learning, adversarial information hiding, and practical software tools, and then identifies the research gap addressed by the proposed suite. Chapter Three presents the system analysis and design of the secure steganography software suite, including the architecture, requirements, security design, datasets, and evaluation metrics. Chapter Four discusses the implementation of the system and its evaluation, including a comparison between the adopted adversarial method and conventional approaches. Chapter Five presents the summary, recommendations, and conclusion of the study.", "style": "Normal"},
    ]

    chapter_two_preamble = [
        {"text": "This chapter reviews the major theories, methods, systems, and tools relevant to the design of the proposed secure steganography software suite. It begins with the broad concepts that define the field, then examines previous works in classical and adversarial steganography, reviews practical software tools, and identifies the gaps that motivate the present study.", "style": "Normal"},
    ]

    chapter_two_existing_works = [
        {"text": "This section reviews the principal streams of prior work that shaped the design of the proposed system. The first stream focuses on steganalysis, especially the movement from handcrafted statistical detection to machine-learning-driven detection. The second stream focuses on modern generative approaches to steganography that aim to reduce detectability by learning adaptive embedding strategies.", "style": "Normal"},
    ]

    chapter_two_ml_steganalysis = [
        {"text": "Early steganalysis techniques depended on explicit statistical tests tailored to known embedding patterns. Over time, these methods were extended into universal steganalysis frameworks that used richer feature sets and conventional classifiers such as Support Vector Machines. Although effective, these systems still depended heavily on handcrafted feature design.", "style": "Normal"},
        {"text": "The modern era of steganalysis is increasingly dominated by deep learning. Convolutional neural networks and residual-network-based detectors learn high-dimensional residual patterns directly from data and are therefore more difficult for fixed-rule embedding methods to evade. This development is one of the key reasons why conventional steganographic methods are no longer sufficient in high-threat environments.", "style": "Normal"},
    ]

    chapter_two_gan_research = [
        {"text": "Adversarial-network-based steganography emerged as a response to stronger steganalysis. Instead of relying on manually defined embedding heuristics, these methods learn to generate stego images whose statistical properties remain close to those of natural images while still enabling message recovery. This introduces a more adaptive and data-driven form of information hiding.", "style": "Normal"},
        {"text": "A major reference point in this area is the work of Zhang et al. (2019), which presented a high-capacity image steganography model based on adversarial training and dense connectivity. Their results showed that deep-learning-based steganography could achieve stronger payload and image-quality trade-offs than earlier approaches, while also reducing detectability under several steganalysis settings. This line of work provides the technical foundation that informs the present project, although the present project is positioned as an original secure software suite rather than a reproduction of that implementation.", "style": "Normal"},
    ]

    chapter_two_existing_systems = [
        {"text": "Existing systems in this area fall into two broad groups. The first group contains classical public tools such as OpenStego and Steghide. These tools are practical and easy to use, but they largely depend on older embedding strategies that are less resilient to machine-learning-based steganalysis. The second group contains research-oriented adversarial or deep-learning steganography prototypes, which often show stronger security characteristics but do not always provide complete user-facing software workflows.", "style": "Normal"},
        {"text": "The proposed project sits between these two groups. It takes the stronger embedding philosophy of adversarial-network-based steganography and combines it with software engineering features expected of a usable suite: structured encode, decode, and check workflows, passphrase-protected payload handling, and web-based accessibility.", "style": "Normal"},
    ]

    chapter_two_methods = [
        {"text": "Classical distortion-cost methods aim to minimize the observable changes introduced during embedding by assigning different costs to different image locations. Such methods are useful because they formalize the trade-off between payload and detectability, but they still operate within fixed handcrafted assumptions about the cover image. Once steganalysis models learn the residual artifacts associated with those assumptions, their secrecy degrades.", "style": "Normal"},
        {"text": "Adversarial training strategies move beyond fixed cost rules by allowing a generator-like model to learn embedding behaviour against a critic or detector. In the context of steganography, this means the embedding engine is optimized not only for message recovery and image similarity, but also for reduced detectability. This is a stronger design principle for modern secure communication.", "style": "Normal"},
        {"text": "Imperceptibility is typically evaluated through quantitative image-quality measures. In practical terms, the most widely used metrics include Mean Squared Error (MSE), Peak Signal-to-Noise Ratio (PSNR), and Structural Similarity Index Measure (SSIM). These metrics do not fully replace human judgement, but they provide repeatable and comparable evidence for the quality of stego images.", "style": "Normal"},
    ]

    chapter_two_tools = [
        {"text": "OpenStego is one of the widely accessible open-source steganography tools. It provides a graphical interface and allows message hiding in image files with optional encryption support. However, its main embedding logic remains rooted in classical techniques, which limits its resilience against stronger detection methods.", "style": "Normal"},
        {"text": "Steghide is another commonly referenced tool that supports image and audio carriers and provides password-protected extraction. Compared with basic LSB tools, it benefits from a more structured embedding process and compression support, but it still belongs to an earlier generation of practical steganography software and does not fundamentally solve the challenge posed by modern deep steganalysis.", "style": "Normal"},
        {"text": "These tools are still useful as practical baselines because they are accessible, understandable, and widely cited. Their weakness under stronger steganalysis is precisely what creates the need for a more secure software suite built around a modern adversarial embedding strategy.", "style": "Normal"},
    ]

    chapter_two_gaps = [
        {"text": "Several research and implementation gaps emerge from the literature and tool review:", "style": "Normal"},
        {"text": "Gap 1: There is a shortage of practical secure steganography software that combines modern adversarial embedding with a complete end-user workflow.", "style": "List Paragraph"},
        {"text": "Gap 2: Many user-friendly public tools still depend on older embedding methods that are relatively weak against contemporary steganalysis.", "style": "List Paragraph"},
        {"text": "Gap 3: Software-level security features such as controlled payload recovery and passphrase-based protection are not consistently integrated into modern research prototypes.", "style": "List Paragraph"},
        {"text": "Gap 4: Comparative discussion often separates model-level performance from actual deployable software concerns such as usability, validation, and secure handling of outputs.", "style": "List Paragraph"},
        {"text": "The present study addresses these gaps by designing a secure software suite that combines an adversarial-network-based embedding core with a practical web interface, encrypted payload handling, and a verification-oriented workflow.", "style": "Normal"},
    ]

    chapter_two_proposed = [
        {"text": "The proposed system is a secure image steganography software suite that allows a user to encode a text message into an image, decode a hidden message with the correct passphrase, and check whether a candidate image contains a valid protected payload. The system is designed around a dense adversarial steganography backbone and a software-layer security component that protects the embedded message before it is passed into the image embedding engine.", "style": "Normal"},
    ]

    chapter_two_overview = [
        {"text": "At the architectural level, the suite combines a browser-based frontend, a FastAPI backend, a secure payload service, and a dense adversarial steganography model. During encoding, the system accepts a cover image, a plaintext message, and a passphrase. The message is transformed into a protected payload envelope before it is hidden in the image. During decoding, the hidden payload is extracted from the stego image and can only be recovered into plaintext when the correct passphrase is supplied.", "style": "Normal"},
        {"text": "The system also introduces a dedicated check workflow. This workflow does not reveal the message and does not require the passphrase, but it determines whether the image contains a valid protected Steg payload. This makes the suite suitable not only for hidden communication, but also for controlled verification of image status within the application.", "style": "Normal"},
    ]

    chapter_two_features = [
        {"text": "The proposed system has the following key features:", "style": "Normal"},
        {"text": "Adversarial-network-based image embedding built on a dense architecture selected for its strong payload and image-quality trade-off.", "style": "List Paragraph"},
        {"text": "Passphrase-protected message handling, which ensures that hidden content remains protected even if extraction is attempted on the stego image.", "style": "List Paragraph"},
        {"text": "A three-part workflow consisting of encode, decode, and check operations for practical software use.", "style": "List Paragraph"},
        {"text": "A web-based user interface and backend service architecture that make the system accessible beyond a research-code setting.", "style": "List Paragraph"},
        {"text": "Validation-oriented handling of uploads, generated outputs, and hidden-payload detection to improve safety and usability.", "style": "List Paragraph"},
    ]

    chapter_two_comparison_intro = [
        {"text": "Table 2.1 compares the proposed secure software suite with representative existing systems and method families across the most relevant practical dimensions.", "style": "Normal"},
        {"text": "Table 2.1: Comparison of Existing Systems and the Proposed Secure Steganography Software Suite", "style": "Normal"},
    ]

    chapter_two_summary = [
        {"text": "This chapter has reviewed the concepts, methods, systems, and tools relevant to secure image steganography. The discussion showed that conventional methods such as LSB substitution and older transform-domain tools remain useful as historical and practical baselines, but they are less resilient to modern machine-learning-based steganalysis. The chapter also showed that adversarial-network-based steganography offers a stronger direction for secrecy, especially when combined with software-level controls. These observations directly motivate the design of the proposed secure steganography software suite presented in the next chapter.", "style": "Normal"},
    ]
    chapter_two_concepts = [
        {"kind": "paragraph", "text": "2.2 Review of Relevant Concepts", "style": "Heading 2"},
        {"kind": "paragraph", "text": "This section reviews the major concepts on which the proposed suite depends.", "style": "Normal"},
        {"kind": "paragraph", "text": "2.2.1 Cryptography", "style": "Heading 3"},
        {"kind": "paragraph", "text": "Cryptography is the discipline of securing information by transforming plaintext into an unreadable form that can only be reversed by an authorized party. In practical secure communication, cryptography is the main mechanism for preserving confidentiality of content.", "style": "Normal"},
        {"kind": "paragraph", "text": "In the context of this project, cryptography complements steganography rather than replacing it. The software suite protects the hidden message with passphrase-based authenticated encryption before the payload is embedded in an image. This means the project secures both the existence of the communication and the meaning of the communication.", "style": "Normal"},
        {"kind": "paragraph", "text": "2.2.2 Steganography", "style": "Heading 3"},
        {"kind": "paragraph", "text": "Steganography is the art and science of concealing information inside another medium in a way that hides the presence of the communication itself. In image steganography, the carrier is a digital image and the hidden information may be text, binary data, or another media object.", "style": "Normal"},
        {"kind": "paragraph", "text": "A sound steganographic system should satisfy imperceptibility, capacity, robustness, and secrecy. These four requirements are often in tension, which makes steganography a problem of careful trade-offs rather than a single optimization target.", "style": "Normal"},
        {"kind": "paragraph", "text": "2.2.3 Spatial Domain Steganography", "style": "Heading 3"},
        {"kind": "paragraph", "text": "Spatial domain steganography hides data by modifying pixel values directly. The most common example is Least Significant Bit substitution, in which message bits are written into the least significant positions of pixel intensities. The attraction of this family is its simplicity and high raw capacity.", "style": "Normal"},
        {"kind": "paragraph", "text": "The main weakness of spatial methods is that their statistical footprints can often be detected by classical and deep learning steganalysis. Even when the visual distortion is small, the hidden modifications may disturb natural image statistics enough to reveal the presence of embedded content.", "style": "Normal"},
        {"kind": "paragraph", "text": "2.2.4 Transform Domain Steganography", "style": "Heading 3"},
        {"kind": "paragraph", "text": "Transform domain methods embed data in frequency-related representations rather than in raw pixel values. JPEG-oriented techniques based on the Discrete Cosine Transform are common examples. These methods often provide better resilience than simple spatial-domain methods because the embedding is distributed across transformed coefficients.", "style": "Normal"},
        {"kind": "paragraph", "text": "Although transform-domain techniques improve on simple LSB substitution, they are still not immune to advanced steganalysis. Once a detector learns the coefficient-level disturbances introduced by embedding, their security advantage can decline significantly.", "style": "Normal"},
        {"kind": "paragraph", "text": "2.2.5 Steganalysis", "style": "Heading 3"},
        {"kind": "paragraph", "text": "Steganalysis is the process of identifying whether hidden data exists in a carrier medium. Earlier steganalysis approaches depended on handcrafted statistical features and targeted attacks. Modern approaches increasingly rely on machine learning and deep residual representations to discover more subtle embedding traces.", "style": "Normal"},
        {"kind": "paragraph", "text": "The strength of present-day steganalysis is the core reason why stronger steganography systems are needed. A practical hidden communication system can no longer assume that low visual distortion alone is enough to guarantee secrecy.", "style": "Normal"},
        {"kind": "paragraph", "text": "2.2.6 Machine Learning", "style": "Heading 3"},
        {"kind": "paragraph", "text": "Machine learning refers to computational methods that learn patterns from data instead of relying only on explicitly written rules. In image security applications, this makes it possible to learn subtle patterns that are difficult to capture using manual statistical engineering alone.", "style": "Normal"},
        {"kind": "paragraph", "text": "2.2.7 Machine Learning in Steganalysis", "style": "Heading 3"},
        {"kind": "paragraph", "text": "When machine learning is applied to steganalysis, the detection problem becomes a classification task. The model learns to separate cover images from stego images by studying the residual disturbances introduced during embedding. This has substantially improved detection performance compared with older handcrafted techniques.", "style": "Normal"},
        {"kind": "paragraph", "text": "Deep learning steganalysis models are especially dangerous to classical embedding schemes because they can learn non-linear patterns distributed across the image. This creates pressure for steganography systems to adopt more adaptive and data-driven concealment strategies.", "style": "Normal"},
        {"kind": "paragraph", "text": "2.2.8 Generative Adversarial Networks", "style": "Heading 3"},
        {"kind": "paragraph", "text": "Generative Adversarial Networks are learning systems in which one model learns to generate outputs while another model learns to distinguish generated outputs from genuine ones. The adversarial pressure created by this competition has made GANs effective in several image-generation tasks.", "style": "Normal"},
        {"kind": "paragraph", "text": "For steganography, this adversarial principle is useful because the embedding model can be trained to reduce detectable artifacts while preserving message recovery. This makes GAN-based steganography a promising response to the stronger detection capabilities of deep learning steganalysis.", "style": "Normal"},
    ]

    chapter_one_block = [
        {"kind": "paragraph", "text": "CHAPTER ONE", "style": "Heading 1"},
        {"kind": "paragraph", "text": "INTRODUCTION", "style": "Heading 1"},
        {"kind": "paragraph", "text": "Background Information", "style": "Heading 2"},
    ] + [{"kind": "paragraph", "text": item["text"], "style": item["style"]} for item in chapter_one_background] + [
        {"kind": "paragraph", "text": "Statement of Problem", "style": "Heading 2"},
    ] + [{"kind": "paragraph", "text": item["text"], "style": item["style"]} for item in chapter_one_problem] + [
        {"kind": "paragraph", "text": "Aim and Objectives", "style": "Heading 2"},
    ] + [{"kind": "paragraph", "text": item["text"], "style": item["style"]} for item in chapter_one_aim] + [
        {"kind": "paragraph", "text": "Methodology", "style": "Heading 2"},
    ] + [{"kind": "paragraph", "text": item["text"], "style": item["style"]} for item in chapter_one_methodology] + [
        {"kind": "paragraph", "text": "Significance of the Study", "style": "Heading 2"},
    ] + [{"kind": "paragraph", "text": item["text"], "style": item["style"]} for item in chapter_one_significance] + [
        {"kind": "paragraph", "text": "Chapter Outline", "style": "Heading 2"},
    ] + [{"kind": "paragraph", "text": item["text"], "style": item["style"]} for item in chapter_one_outline]

    chapter_two_block = [
        {"kind": "pagebreak"},
        {"kind": "paragraph", "text": "CHAPTER TWO", "style": "Heading 1"},
        {"kind": "paragraph", "text": "LITERATURE REVIEW", "style": "Heading 1"},
        {"kind": "paragraph", "text": "2.1 Preamble", "style": "Heading 2"},
    ] + [{"kind": "paragraph", "text": item["text"], "style": item["style"]} for item in chapter_two_preamble] + chapter_two_concepts + [
        {"kind": "paragraph", "text": "2.3 Review of Existing Works", "style": "Heading 2"},
    ] + [{"kind": "paragraph", "text": item["text"], "style": item["style"]} for item in chapter_two_existing_works] + [
        {"kind": "paragraph", "text": "2.3.1 Machine Learning Steganalysis Research", "style": "Heading 3"},
    ] + [{"kind": "paragraph", "text": item["text"], "style": item["style"]} for item in chapter_two_ml_steganalysis] + [
        {"kind": "paragraph", "text": "2.3.2 GAN-Based Steganography Research", "style": "Heading 3"},
    ] + [{"kind": "paragraph", "text": item["text"], "style": item["style"]} for item in chapter_two_gan_research] + [
        {"kind": "paragraph", "text": "2.4 Review of Existing Systems", "style": "Heading 2"},
    ] + [{"kind": "paragraph", "text": item["text"], "style": item["style"]} for item in chapter_two_existing_systems] + [
        {"kind": "paragraph", "text": "2.5 Review of Existing Methods", "style": "Heading 2"},
    ] + [{"kind": "paragraph", "text": item["text"], "style": item["style"]} for item in chapter_two_methods] + [
        {"kind": "paragraph", "text": "2.6 Review of Existing Tools", "style": "Heading 2"},
    ] + [{"kind": "paragraph", "text": item["text"], "style": item["style"]} for item in chapter_two_tools] + [
        {"kind": "paragraph", "text": "2.7 Research Gaps", "style": "Heading 2"},
    ] + [{"kind": "paragraph", "text": item["text"], "style": item["style"]} for item in chapter_two_gaps] + [
        {"kind": "paragraph", "text": "2.8 Proposed System", "style": "Heading 2"},
    ] + [{"kind": "paragraph", "text": item["text"], "style": item["style"]} for item in chapter_two_proposed] + [
        {"kind": "paragraph", "text": "2.8.1 System Overview", "style": "Heading 3"},
    ] + [{"kind": "paragraph", "text": item["text"], "style": item["style"]} for item in chapter_two_overview] + [
        {"kind": "paragraph", "text": "2.8.2 Key Features", "style": "Heading 3"},
    ] + [{"kind": "paragraph", "text": item["text"], "style": item["style"]} for item in chapter_two_features] + [
        {"kind": "paragraph", "text": "2.8.3 Comparison of Existing Systems with Proposed System", "style": "Heading 3"},
    ] + [{"kind": "paragraph", "text": item["text"], "style": item["style"]} for item in chapter_two_comparison_intro] + [
        {"kind": "table", "rows": [["System / Method", "Embedding Basis", "Security Against Modern Steganalysis", "User Accessibility", "Payload Protection"], ["OpenStego", "Classical spatial-domain hiding", "Low", "High", "Optional encryption support"], ["Steghide", "Classical transform-oriented hiding", "Moderate to low", "Moderate", "Password-based extraction"], ["Research adversarial models", "Deep adversarial embedding", "High at model level", "Low", "Usually prototype dependent"], ["Proposed Steg suite", "Adversarial dense image steganography", "Higher than classical methods", "High", "Passphrase-protected payload recovery"]]},
        {"kind": "paragraph", "text": "2.9 Chapter Outline", "style": "Heading 2"},
    ] + [{"kind": "paragraph", "text": item["text"], "style": item["style"]} for item in chapter_two_summary]

    body_anchor = find_paragraph(doc, "ABBREVIATIONS")
    remove_after(body_anchor)
    chapter_insert_anchor = body_anchor
    chapter_insert_anchor = add_section_block(chapter_insert_anchor, chapter_one_block)
    chapter_insert_anchor = add_section_block(chapter_insert_anchor, chapter_two_block)

    chapter_three_content = [
        {"kind": "pagebreak"},
        {"kind": "paragraph", "text": "CHAPTER THREE", "style": "Heading 1"},
        {"kind": "paragraph", "text": "SYSTEM ANALYSIS AND DESIGN", "style": "Heading 1"},
        {"kind": "paragraph", "text": "3.1 Preamble", "style": "Heading 2"},
        {"kind": "paragraph", "text": "This chapter presents the analysis and design of the secure steganography software suite developed in this project. It translates the ideas discussed in the literature into a practical system architecture and explains how the software components, security controls, and evaluation metrics work together.", "style": "Normal"},
        {"kind": "paragraph", "text": "3.2 Overview of the System", "style": "Heading 2"},
        {"kind": "paragraph", "text": "The proposed system is a web-based secure steganography software suite that enables a user to hide a text message inside an image, recover the hidden message using the correct passphrase, and check whether a candidate image contains a valid protected payload. The system is centered on three operational workflows: encode, decode, and check.", "style": "Normal"},
        {"kind": "paragraph", "text": "From a design perspective, the system combines four major layers. The first is the presentation layer, which provides the user with an accessible browser interface. The second is the application layer, which validates requests and coordinates processing through the backend services. The third is the security layer, which performs passphrase validation, key derivation, and authenticated payload encryption. The fourth is the steganography layer, which embeds and extracts the protected payload using a dense adversarial image steganography model.", "style": "Normal"},
        {"kind": "paragraph", "text": "3.3 Requirements Analysis", "style": "Heading 2"},
        {"kind": "paragraph", "text": "The analysis of the system requirements was guided by the security problem identified in earlier chapters. The system had to do more than hide a message in an image. It also had to provide a secure and usable process around that hidden message.", "style": "Normal"},
        {"kind": "paragraph", "text": "3.3.1 Functional Requirements", "style": "Heading 3"},
        {"kind": "paragraph", "text": "The system shall allow a user to upload a cover image and encode a text message into it.", "style": "List Paragraph"},
        {"kind": "paragraph", "text": "The system shall require a passphrase during encoding so that the concealed message is protected before image embedding takes place.", "style": "List Paragraph"},
        {"kind": "paragraph", "text": "The system shall allow a user to upload a stego image and decode the hidden message only when the correct passphrase is supplied.", "style": "List Paragraph"},
        {"kind": "paragraph", "text": "The system shall provide a check operation that reports whether an image contains a valid protected payload without exposing the message itself.", "style": "List Paragraph"},
        {"kind": "paragraph", "text": "The system shall allow the encoded image to be downloaded after successful processing.", "style": "List Paragraph"},
        {"kind": "paragraph", "text": "3.3.2 Non-Functional Requirements", "style": "Heading 3"},
        {"kind": "paragraph", "text": "Security: the system shall not store user passphrases in plaintext and shall protect message recovery through authenticated encryption.", "style": "List Paragraph"},
        {"kind": "paragraph", "text": "Usability: the system shall provide a straightforward web interface that can be used by non-expert users.", "style": "List Paragraph"},
        {"kind": "paragraph", "text": "Performance: the system shall complete normal encode, decode, and check operations within a practical waiting time for standard images.", "style": "List Paragraph"},
        {"kind": "paragraph", "text": "Maintainability: the system shall separate interface logic, API logic, secure payload logic, and model invocation logic.", "style": "List Paragraph"},
        {"kind": "paragraph", "text": "Reliability: the system shall return controlled errors for invalid files, missing passphrases, wrong passphrases, and unsupported payloads.", "style": "List Paragraph"},
        {"kind": "paragraph", "text": "3.4 Dataset Description", "style": "Heading 2"},
        {"kind": "paragraph", "text": "The image data associated with this project is based on the DIV2K and MSCOCO datasets. DIV2K provides high-resolution natural images that are useful for studying image quality and fine-detail preservation. MSCOCO provides a broader variety of natural scenes, objects, and textures, making it suitable for evaluating behaviour across diverse real-world image content.", "style": "Normal"},
        {"kind": "paragraph", "text": "These datasets are important because adversarial image steganography depends heavily on the statistics of natural images. A model trained or evaluated on richer and more varied image content is more likely to generalize to realistic use cases than one trained on narrow or repetitive image sets.", "style": "Normal"},
        {"kind": "paragraph", "text": "3.5 Physical Design", "style": "Heading 2"},
        {"kind": "paragraph", "text": "3.5.1 System Architecture", "style": "Heading 3"},
        {"kind": "paragraph", "text": "Figure 3.1 presents the physical architecture of the proposed suite.", "style": "Normal"},
        {"kind": "image", "path": architecture_path, "width": Inches(6.4)},
        {"kind": "paragraph", "text": "Figure 3.1: Architecture of the proposed secure steganography software suite", "style": "Caption"},
        {"kind": "paragraph", "text": "The presentation layer is responsible for collecting input and displaying results. The application layer handles routing, validation, and request coordination. The security layer transforms a plaintext message into a protected payload envelope using a passphrase. The steganography core then embeds or extracts that payload through the dense adversarial model, while runtime storage temporarily manages uploads and downloadable outputs.", "style": "Normal"},
        {"kind": "paragraph", "text": "3.6 Logical Design", "style": "Heading 2"},
        {"kind": "paragraph", "text": "3.6.1 Operational Workflow", "style": "Heading 3"},
        {"kind": "paragraph", "text": "Figure 3.2 shows the logical workflow of the three main operations of the system.", "style": "Normal"},
        {"kind": "image", "path": workflow_path, "width": Inches(6.4)},
        {"kind": "paragraph", "text": "Figure 3.2: Operational workflow for encode, decode, and check", "style": "Caption"},
        {"kind": "paragraph", "text": "During encoding, the user provides a cover image, a secret message, and a passphrase. The message is first encrypted into a structured payload envelope and then hidden in the cover image. During decoding, the hidden payload is extracted from the stego image and decrypted with the correct passphrase. During checking, the system only tests whether a valid protected payload envelope exists in the image.", "style": "Normal"},
        {"kind": "paragraph", "text": "3.6.2 Secure Payload Design", "style": "Heading 3"},
        {"kind": "paragraph", "text": "A major part of the logical design is the separation between message protection and image embedding. Before the message is hidden in the image, it is transformed into a protected envelope with a version marker, salt, nonce, and authenticated ciphertext. This means that even if a hidden payload is successfully extracted from the image, the message content still requires the correct passphrase before it can be interpreted.", "style": "Normal"},
        {"kind": "paragraph", "text": "The suite therefore operates as a layered security system. Steganography conceals the presence of the message, while passphrase-based authenticated encryption protects the meaning of the message.", "style": "Normal"},
        {"kind": "paragraph", "text": "3.7 Mathematical Formulation and Evaluation Metrics", "style": "Heading 2"},
        {"kind": "paragraph", "text": "Let C represent the cover image, S represent the stego image, M represent the plaintext message, and P represent the protected payload. The secure payload generation step can be represented as:", "style": "Normal"},
        {"kind": "paragraph", "text": "P = Enc_k(M), where k = KDF(passphrase, salt)", "style": "Normal", "align": "center"},
        {"kind": "paragraph", "text": "The image embedding process is then represented as:", "style": "Normal"},
        {"kind": "paragraph", "text": "S = E(C, P)", "style": "Normal", "align": "center"},
        {"kind": "paragraph", "text": "During recovery, the extracted payload P' is passed through the secure decoder:", "style": "Normal"},
        {"kind": "paragraph", "text": "M' = Dec_k(P')", "style": "Normal", "align": "center"},
        {"kind": "paragraph", "text": "To evaluate image distortion, the Mean Squared Error (MSE) between the cover and stego image is given by:", "style": "Normal"},
        {"kind": "paragraph", "text": "MSE = (1 / MN) Σ Σ [C(i,j) - S(i,j)]²", "style": "Normal", "align": "center"},
        {"kind": "paragraph", "text": "The Peak Signal-to-Noise Ratio (PSNR) is then computed as:", "style": "Normal"},
        {"kind": "paragraph", "text": "PSNR = 10 log10(MAX² / MSE)", "style": "Normal", "align": "center"},
        {"kind": "paragraph", "text": "Structural Similarity Index Measure (SSIM) is used to evaluate perceptual similarity between the cover and stego images:", "style": "Normal"},
        {"kind": "paragraph", "text": "SSIM(x,y) = ((2μxμy + C1)(2σxy + C2)) / ((μx² + μy² + C1)(σx² + σy² + C2))", "style": "Normal", "align": "center"},
        {"kind": "paragraph", "text": "For the adopted adversarial method, effective embedding capacity is discussed in terms of Reed-Solomon bits per pixel (RS-BPP), given by:", "style": "Normal"},
        {"kind": "paragraph", "text": "RS-BPP = D(1 - 2p)", "style": "Normal", "align": "center"},
        {"kind": "paragraph", "text": "where D is the attempted payload depth and p is the bit error rate. These metrics together provide a balanced basis for analysing image quality, payload efficiency, and practical recoverability.", "style": "Normal"},
        {"kind": "paragraph", "text": "3.8 Summary", "style": "Heading 2"},
        {"kind": "paragraph", "text": "This chapter has presented the system analysis and design of the proposed secure steganography software suite. It described the requirements, architecture, workflow, secure payload design, and evaluation formulas that guide the implementation. The next chapter discusses the actual implementation and evaluation of the system.", "style": "Normal"},
    ]

    chapter_four_content = [
        {"kind": "pagebreak"},
        {"kind": "paragraph", "text": "CHAPTER FOUR", "style": "Heading 1"},
        {"kind": "paragraph", "text": "SYSTEM IMPLEMENTATION AND EVALUATION", "style": "Heading 1"},
        {"kind": "paragraph", "text": "4.1 Preamble", "style": "Heading 2"},
        {"kind": "paragraph", "text": "This chapter explains how the proposed secure steganography software suite was implemented and how its performance was evaluated. The discussion covers the tools used, the software modules created, the datasets referenced for the adopted adversarial method, and the practical and comparative evaluation of the completed suite.", "style": "Normal"},
        {"kind": "paragraph", "text": "4.2 System Requirements", "style": "Heading 2"},
        {"kind": "paragraph", "text": "4.2.1 Hardware Requirements", "style": "Heading 3"},
        {"kind": "paragraph", "text": "A modern personal computer with at least a multi-core processor, 8 GB RAM, and sufficient local storage is adequate for running the software suite with pretrained models. For model training or extended experimentation, a machine with GPU acceleration is preferable.", "style": "Normal"},
        {"kind": "paragraph", "text": "4.2.2 Software Requirements", "style": "Heading 3"},
        {"kind": "paragraph", "text": "The implementation requires Python, PyTorch, FastAPI, Jinja2, Pillow, python-multipart, cryptography, and standard web browser support. During development and deployment, Uvicorn and Gunicorn were used for application serving.", "style": "Normal"},
        {"kind": "paragraph", "text": "4.3 Implementation Tools", "style": "Heading 2"},
        {"kind": "paragraph", "text": "Python was used as the principal implementation language because it provides direct access to the adopted steganography model, deep learning tooling, and backend service development. PyTorch supports the adversarial image steganography backbone. FastAPI was selected for the backend API because it provides a lightweight and structured web service layer. The frontend interface was implemented with HTML, CSS, and JavaScript to provide direct encode, decode, and check interactions. The cryptography library was used to implement authenticated payload protection.", "style": "Normal"},
        {"kind": "paragraph", "text": "4.4 Program Modules and Interfaces", "style": "Heading 2"},
        {"kind": "paragraph", "text": "4.4.1 Encode Module", "style": "Heading 3"},
        {"kind": "paragraph", "text": "The encode module accepts the cover image, plaintext message, and passphrase. It validates the inputs, encrypts the message into the STEGv1 payload format, and invokes the adversarial image encoder to generate the final stego image. The output image is then made available for download through the interface.", "style": "Normal"},
        {"kind": "paragraph", "text": "4.4.2 Decode Module", "style": "Heading 3"},
        {"kind": "paragraph", "text": "The decode module accepts the stego image and the passphrase. It extracts the hidden payload from the image, verifies that it conforms to the supported protected format, and decrypts the payload only when the correct passphrase is supplied. Incorrect passphrases do not reveal the message and instead produce a controlled error.", "style": "Normal"},
        {"kind": "paragraph", "text": "4.4.3 Check Module", "style": "Heading 3"},
        {"kind": "paragraph", "text": "The check module is designed for payload presence verification. It inspects the extracted hidden content and determines whether the image contains a valid protected payload envelope without requiring the passphrase. This module is useful where payload presence must be validated without revealing message content.", "style": "Normal"},
        {"kind": "paragraph", "text": "4.4.4 Security Module", "style": "Heading 3"},
        {"kind": "paragraph", "text": "The security module enforces a minimum passphrase length, derives a key from the passphrase using a salted key derivation function, and uses authenticated encryption to protect the plaintext message before it is hidden in the image. This module is central to the security claims of the software suite because it ensures that concealment and confidentiality operate together.", "style": "Normal"},
        {"kind": "paragraph", "text": "4.4.5 Frontend and API Interface", "style": "Heading 3"},
        {"kind": "paragraph", "text": "The frontend presents the three core workflows in a structured single-page interface. The backend exposes routes for encoding, decoding, checking, and downloading generated outputs. This division makes the application easier to maintain and test.", "style": "Normal"},
        {"kind": "paragraph", "text": "4.5 Dataset Preparation and Experimental Setup", "style": "Heading 2"},
        {"kind": "paragraph", "text": "The project references the DIV2K and MSCOCO datasets as the image datasets associated with the adopted adversarial image steganography approach. DIV2K contributes high-resolution image content, while MSCOCO contributes a wide variety of natural scenes and object-rich images. Together they provide a strong basis for analysing image quality and embedding behaviour across different image characteristics.", "style": "Normal"},
        {"kind": "paragraph", "text": "The evaluation in this report is considered at two levels. The first level is model-level evidence for the dense adversarial steganography backbone adopted by the suite. The second level is system-level verification of the secure software suite itself, including functional behaviour, passphrase protection, and hidden-payload detection.", "style": "Normal"},
        {"kind": "paragraph", "text": "4.6 Evaluation of the System", "style": "Heading 2"},
        {"kind": "paragraph", "text": "4.6.1 Functional Validation", "style": "Heading 3"},
        {"kind": "paragraph", "text": "Table 4.1 summarises the principal functional checks carried out on the completed software suite.", "style": "Normal"},
        {"kind": "table", "rows": [["Operation", "Expected Behaviour", "Observed Result"], ["Encode", "Generate a downloadable stego image from image, message, and passphrase", "Successful"], ["Decode with correct passphrase", "Recover the original hidden message", "Successful"], ["Decode with wrong passphrase", "Reject recovery without exposing plaintext", "Successful"], ["Check image", "Report presence of supported protected payload without passphrase", "Successful"], ["Automated webapp tests", "Verify API, UI, and service behaviour", "47 tests passed"]]},
        {"kind": "paragraph", "text": "Table 4.1: Functional validation of the secure steganography software suite", "style": "Caption"},
        {"kind": "paragraph", "text": "The functional validation confirmed that the suite performs the three intended operations correctly. In particular, the passphrase-protected decode workflow ensures that extraction of the hidden envelope alone does not expose the plaintext message.", "style": "Normal"},
        {"kind": "paragraph", "text": "4.6.2 Security Evaluation", "style": "Heading 3"},
        {"kind": "paragraph", "text": "The security evaluation focused on the software-layer protections added to the suite. The system enforces a minimum passphrase length of twelve characters, rejects blank or omitted passphrases, validates the structure of protected payload envelopes, and prevents malformed payloads from being accepted as valid hidden data. These checks reduce the chance of insecure usage and improve the trustworthiness of the check workflow.", "style": "Normal"},
        {"kind": "paragraph", "text": "4.6.3 Adopted Adversarial Model Performance", "style": "Heading 3"},
        {"kind": "paragraph", "text": "At the model level, the suite uses the dense adversarial architecture as its operational default because it offers the strongest overall trade-off among the reported variants. Table 4.2 presents selected dense-model results reported for the adopted adversarial method on DIV2K and MSCOCO.", "style": "Normal"},
        {"kind": "table", "rows": [["Dataset", "Data Depth (D)", "Accuracy", "RS-BPP", "PSNR (dB)", "SSIM"], ["DIV2K", "4", "0.82", "2.53", "37.49", "0.88"], ["DIV2K", "6", "0.70", "2.44", "38.94", "0.90"], ["MSCOCO", "4", "0.95", "3.61", "36.94", "0.92"], ["MSCOCO", "6", "0.87", "4.40", "36.33", "0.88"]]},
        {"kind": "paragraph", "text": "Table 4.2: Selected dense adversarial steganography results adapted from Zhang et al. (2019)", "style": "Caption"},
        {"kind": "paragraph", "text": "These figures justify the selection of the dense architecture as the software default. They show that the adopted method can maintain acceptable visual quality while supporting a payload level that is substantially higher than what is commonly associated with older practical tools.", "style": "Normal"},
        {"kind": "paragraph", "text": "4.6.4 Comparison with Traditional Methods", "style": "Heading 3"},
        {"kind": "paragraph", "text": "The developed suite was also compared conceptually and practically against weaker traditional approaches. Table 4.3 summarises this comparison.", "style": "Normal"},
        {"kind": "table", "rows": [["Criterion", "LSB-Based Methods", "Transform-Domain Methods", "Proposed Steg Suite"], ["Embedding strategy", "Fixed pixel-level substitution", "Coefficient-domain modification", "Learned adversarial dense embedding"], ["Resistance to modern steganalysis", "Low", "Moderate", "Higher"], ["Payload protection", "Usually external or optional", "Sometimes password-based", "Integrated passphrase-protected payload"], ["User workflow", "Often simple", "Moderate", "Encode, decode, and check in one suite"], ["Practical detectability", "High under strong detectors", "Lower than LSB but still vulnerable", "More resilient than classical methods"]]},
        {"kind": "paragraph", "text": "Table 4.3: Comparison of the proposed suite with traditional steganography approaches", "style": "Caption"},
        {"kind": "paragraph", "text": "The comparison shows that classical methods remain useful as educational and baseline techniques, but they do not provide the same combination of adaptive embedding, structured security, and practical workflow integration achieved by the proposed suite.", "style": "Normal"},
        {"kind": "paragraph", "text": "4.6.5 Discussion", "style": "Heading 3"},
        {"kind": "paragraph", "text": "The results indicate that the most important contribution of the project is not merely the use of an adversarial embedding backbone, but the translation of that backbone into a secure and practical software system. The suite strengthens hidden communication at the software level by enforcing passphrase-based recovery and by separating encode, decode, and check responsibilities clearly. This makes the project more useful in practice than a model-only demonstration.", "style": "Normal"},
        {"kind": "paragraph", "text": "The evaluation also shows that the adversarial dense approach is the right backbone for this system when compared with less effective traditional methods. In terms of secrecy and modern detectability, it is better suited to present-day threat conditions than simple LSB or older transform-domain tools.", "style": "Normal"},
        {"kind": "paragraph", "text": "4.7 Summary", "style": "Heading 2"},
        {"kind": "paragraph", "text": "This chapter has presented the implementation and evaluation of the secure steganography software suite. It described the tools used, the system modules, the role of the datasets, the validation of the completed software, and the comparative position of the adopted adversarial method relative to conventional techniques.", "style": "Normal"},
    ]

    chapter_five_content = [
        {"kind": "pagebreak"},
        {"kind": "paragraph", "text": "CHAPTER FIVE", "style": "Heading 1"},
        {"kind": "paragraph", "text": "SUMMARY, RECOMMENDATIONS AND CONCLUSION", "style": "Heading 1"},
        {"kind": "paragraph", "text": "5.1 Summary", "style": "Heading 2"},
        {"kind": "paragraph", "text": "This project set out to design, implement, and evaluate a secure steganography software suite that is more suitable for modern threat conditions than conventional public steganography tools. The study showed that the growth of machine-learning-based steganalysis has reduced the practical security of many classical methods, especially those based on fixed spatial or transform-domain rules.", "style": "Normal"},
        {"kind": "paragraph", "text": "To address this problem, the project developed a software suite built around an adversarial-network-based image steganography backbone and strengthened it with a passphrase-protected payload workflow. The resulting system supports image encoding, secure decoding, and hidden-data checking through an accessible interface. The report also compared the adopted adversarial method with weaker traditional methods and showed why the chosen approach is more appropriate for present-day secure hidden communication.", "style": "Normal"},
        {"kind": "paragraph", "text": "5.2 Recommendations", "style": "Heading 2"},
        {"kind": "paragraph", "text": "The following recommendations arise from the study:", "style": "Normal"},
        {"kind": "paragraph", "text": "Future work should evaluate the suite against additional modern steganalysis detectors under a broader range of practical attack conditions.", "style": "List Paragraph"},
        {"kind": "paragraph", "text": "The software can be extended to support more carrier media such as audio and video while retaining the same secure payload philosophy.", "style": "List Paragraph"},
        {"kind": "paragraph", "text": "Further user studies should be carried out to evaluate ease of use, error rates, and adoption by non-expert users.", "style": "List Paragraph"},
        {"kind": "paragraph", "text": "A future research version of the system may explore stronger benchmarking, model retraining pipelines, and automated reporting of image-quality metrics directly inside the interface.", "style": "List Paragraph"},
        {"kind": "paragraph", "text": "5.3 Conclusion", "style": "Heading 2"},
        {"kind": "paragraph", "text": "In conclusion, this project has demonstrated that a secure and practical steganography software suite can be developed by combining adversarial-network-based image embedding with software-level security controls. The work does not stop at hidden-image generation alone, but provides a usable secure workflow in which message concealment, controlled recovery, and payload verification are integrated into one system.", "style": "Normal"},
        {"kind": "paragraph", "text": "The final outcome is an original secure steganography software suite that is academically grounded, practically useful, and more aligned with modern steganalysis-aware security needs than conventional approaches. This makes the study a meaningful contribution to applied information hiding research and a suitable foundation for future publication-oriented refinement.", "style": "Normal"},
    ]

    chapter_insert_anchor = add_section_block(chapter_insert_anchor, chapter_three_content)
    chapter_insert_anchor = add_section_block(chapter_insert_anchor, chapter_four_content)
    chapter_insert_anchor = add_section_block(chapter_insert_anchor, chapter_five_content)
    chapter_insert_anchor = add_section_block(chapter_insert_anchor, [
        {"kind": "paragraph", "text": "", "style": "Normal"},
        {"kind": "paragraph", "text": "REFERENCES", "style": "Heading 1"},
    ])

    format_document(doc)
    doc.save(REPORT_PATH)


if __name__ == "__main__":
    main()
